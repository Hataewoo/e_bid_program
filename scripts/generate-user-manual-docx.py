# -*- coding: utf-8 -*-
"""Generate CS E-Bid Analyzer user manual (DOCX)."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn

OUT = Path(__file__).resolve().parents[1] / "docs" / "CS-E-Bid-Analyzer-사용법-v1.0.0.docx"


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")


def add_title(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_p(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
    doc.add_paragraph()


def build() -> None:
    doc = Document()
    set_doc_defaults(doc)

    sections = doc.sections[0]
    sections.top_margin = Cm(2.5)
    sections.bottom_margin = Cm(2.5)
    sections.left_margin = Cm(2.5)
    sections.right_margin = Cm(2.5)

    add_title(doc, "CS E-Bid Analyzer\n상세 사용 설명서")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"버전 1.0.0  |  작성일 {date.today().isoformat()}\n")
    meta.add_run("Windows 전자입찰 분석 프로그램 (Electron Desktop)")

    doc.add_page_break()

    # 1
    add_h(doc, "1. 개요")
    add_p(
        doc,
        "CS E-Bid Analyzer는 레거시 MFC 프로그램 「전자입찰 누적카운트」의 핵심 기능을 "
        "현대 Windows 데스크톱 환경에서 제공하는 분석 도구입니다. Master(00~99) 데이터, "
        "Code 패턴, STEP2/3 분석, 통계, 역분석(Research Lab)을 하나의 앱에서 관리합니다.",
    )
    add_bullets(
        doc,
        [
            "대상 OS: Windows 10/11 (64비트)",
            "데이터 저장: 로컬 SQLite (단일 사용자)",
            "언어: 한국어 / English (Settings에서 전환)",
            "릴리스: v1.0.0 (2026-06-30)",
        ],
    )

    # 2
    add_h(doc, "2. 설치 및 첫 실행")
    add_h(doc, "2.1 설치 (배포본)", 2)
    add_numbered(
        doc,
        [
            "배포 파일 `CS E-Bid Analyzer-Setup-1.0.0-x64.exe`를 실행합니다.",
            "설치 마법사에서 경로를 확인한 뒤 설치를 완료합니다.",
            "바탕화면 또는 시작 메뉴의 「CS E-Bid Analyzer」로 앱을 실행합니다.",
            "최초 실행 시 사용자 데이터 폴더에 SQLite DB가 자동 생성됩니다.",
        ],
    )
    add_p(doc, "DB 기본 경로 (패키징 설치본):", bold=True)
    add_p(doc, "%APPDATA%\\CS E-Bid Analyzer\\database\\cs_e_bid.db")

    add_h(doc, "2.2 개발자 실행 (소스에서)", 2)
    add_numbered(
        doc,
        [
            "Node.js 20 이상을 설치합니다.",
            "프로젝트 폴더에서 `npm install` 후 `npm run dev`로 실행합니다.",
            "품질 검증: `npm run build:check` (lint + 테스트 + 회귀 게이트 + 빌드).",
        ],
    )

    # 3
    add_h(doc, "3. 화면 구성")
    add_p(doc, "앱은 Windows 클래식 스타일 레이아웃으로 구성됩니다.")
    add_table(
        doc,
        ["영역", "설명"],
        [
            ["상단 Toolbar", "앱 이름, DB 연결 상태, 테마 전환, 일괄 가져오기, 프로그램 정보"],
            ["좌측 Sidebar", "MASTER, CODE, Reverse Engineering, Research, Analysis, Statistics, Settings 메뉴"],
            ["중앙 작업 영역", "선택한 메뉴의 기능 화면"],
            ["하단 StatusBar", "저장/오류/진행 상태 메시지"],
        ],
    )
    add_h(doc, "3.1 메뉴 순서 변경 (Menu Edit)", 2)
    add_bullets(
        doc,
        [
            "사이드바 각 항목 왼쪽의 ⋮⋮ 핸들을 드래그하여 순서를 변경할 수 있습니다.",
            "↺ 버튼으로 메뉴 순서를 초기화합니다.",
            "« / » 버튼으로 사이드바를 접거나 펼칩니다.",
            "v1.0에서는 사이드바 순서 변경만 지원합니다 (전체 레거시 메뉴 편집기는 미포함).",
        ],
    )

    # 4
    add_h(doc, "4. MASTER — 마스터 데이터 (00~99)")
    add_p(doc, "입찰 분석의 기준이 되는 숫자열(Master Value)을 슬롯 00~99에 저장합니다.")
    add_h(doc, "4.1 기본 작업", 2)
    add_numbered(
        doc,
        [
            "Sidebar에서 MASTER를 선택합니다.",
            "좌측 그리드에서 슬롯(00~99)을 클릭하거나 콤보박스로 번호를 선택합니다.",
            "우측 편집기에 Master Value(숫자만, 최대 1,000자)와 Memo를 입력합니다.",
            "저장 버튼 또는 Ctrl+S로 저장합니다.",
            "신규(F2), 삭제(Delete)로 레코드를 관리합니다.",
        ],
    )
    add_h(doc, "4.2 검증", 2)
    add_bullets(
        doc,
        [
            "Master 번호는 00~99만 허용됩니다.",
            "Value는 숫자만 입력 가능합니다 (공백·쉼표는 저장 시 정리).",
            "「검증」 버튼으로 형식 오류를 사전 확인할 수 있습니다.",
        ],
    )
    add_h(doc, "4.3 가져오기 /보내기", 2)
    add_bullets(
        doc,
        [
            "화면 내 CSV·Excel 가져오기/보내기 지원.",
            "Toolbar 「데이터 일괄 가져오기」(Ctrl+Shift+I): TXT/CSV/XLSX로 Master 00~99 대량 등록.",
            "TXT 형식: 마스터번호,숫자배열,비고 / CSV: masterNo, value, memo",
        ],
    )

    # 5
    add_h(doc, "5. CODE — 패턴 코드 관리")
    add_numbered(
        doc,
        [
            "Sidebar에서 CODE를 선택합니다.",
            "검색창으로 코드명·설명을 필터링합니다.",
            "그리드에서 행을 선택하고 Code, Type, Description을 편집합니다.",
            "F2(신규), Ctrl+S(저장), Delete(삭제) 단축키를 사용할 수 있습니다.",
            "CSV·Excel로 일괄 가져오기/보내기가 가능합니다.",
        ],
    )

    # 6
    add_h(doc, "6. Code Value — 페이지별 코드 카운트")
    add_p(
        doc,
        "Code 기반 Value 분석 UI입니다. 패턴 통계를 확인할 수 있으나, "
        "v1.0에서는 레거시 1:1 검증이 완료되지 않았음을 나타내는 배너가 표시됩니다.",
    )
    add_bullets(
        doc,
        [
            "Sidebar에서 CodeValue 메뉴(해시 라우트 /code-value)로 이동합니다.",
            "Master·Code 선택 후 분석 결과를 확인합니다.",
            "결과 해석 시 「레거시 미검증」 안내를 참고하세요.",
        ],
    )

    # 7
    add_h(doc, "7. Reverse Engineering")
    add_p(doc, "Master Value의 구조를 STEP1~6 패널로 분석합니다 (클라이언트 측).")
    add_table(
        doc,
        ["단계", "내용"],
        [
            ["STEP1", "숫자 빈도 분석"],
            ["STEP2", "Low digit(0~4) 추출"],
            ["STEP3", "High digit(5~9) 추출"],
            ["STEP4", "연속 그룹 분석"],
            ["STEP5", "Run-Length Encoding"],
            ["STEP6", "JSON 형태 출력"],
        ],
    )
    add_p(doc, "결과는 JSON/TXT/CSV로 복사·보내기할 수 있으며, Research 실험 입력 참고 자료로 활용합니다.")

    # 8
    add_h(doc, "8. Analysis — 입찰 분석")
    add_h(doc, "8.1 단일 분석", 2)
    add_numbered(
        doc,
        [
            "Analysis 메뉴로 이동합니다.",
            "Master와 Code를 선택합니다.",
            "「Load」 후 「Analyze」로 STEP2/3, Statistics, CodeValue 통계, Prediction을 실행합니다.",
            "분석 결과는 화면 패널과 Raw JSON 뷰에서 확인합니다.",
            "JSON보내기로 결과를 파일로 저장할 수 있습니다.",
        ],
    )
    add_h(doc, "8.2 배치 분석", 2)
    add_p(doc, "등록된 Master 00~99 슬롯을 순차 분석하는 일괄 분석 모달을 제공합니다. 진행률과 CSV보내기를 지원합니다.")

    add_h(doc, "8.3 분석 이력", 2)
    add_p(doc, "이력 패널에서 과거 분석 기록을 조회·삭제할 수 있습니다. DB 저장 실패 시 하단 StatusBar에 알림이 표시됩니다.")

    add_h(doc, "8.4 Prediction (예측)", 2)
    add_p(
        doc,
        "규칙 기반 휴리스틱으로 동작합니다. 레거시 프로그램과 1:1 일치가 보장되지 않으며, "
        "화면에 미검증 안내 배너가 표시됩니다.",
    )

    # 9
    add_h(doc, "9. Statistics — 통계")
    add_bullets(
        doc,
        [
            "Frequency: 숫자별 출현 빈도",
            "Low/High Ratio: 저자리·고자리 비율",
            "Distribution: 분포 차트",
            "분석 이력 기반 집계 및 CSV/Excel보내기",
        ],
    )

    # 10
    add_h(doc, "10. Algorithm Research (Research Lab)")
    add_p(doc, "역분석 실험·가설·검증(Test Case)·Test Suite를 관리하는 연구 워크스페이스입니다.")
    add_h(doc, "10.1 주요 탭", 2)
    add_table(
        doc,
        ["탭", "용도"],
        [
            ["Experiments", "레거시 vs 신규 프로그램 출력 비교 실험"],
            ["Hypotheses", "알고리즘 가설 등록·추적"],
            ["Verifications / Test Cases", "Pass/Fail 검증 케이스"],
            ["Test Suite", "내장 회귀·카탈로그 진단 실행"],
            ["Dashboard", "통과율 추이, 최근 FAIL 목록"],
        ],
    )
    add_h(doc, "10.2 권장 워크플로", 2)
    add_numbered(
        doc,
        [
            "레거시 프로그램에서 관측값(Expected)을 기록합니다.",
            "Experiment를 생성하고 입력·출력을 저장합니다.",
            "신규 프로그램 결과(Actual)와 비교합니다.",
            "차이가 있으면 Hypothesis를 작성하고 Test Case로 등록합니다.",
            "Test Suite에서 built-in regression 또는 catalog diagnose를 실행합니다.",
        ],
    )
    add_p(doc, "카탈로그 JSON/CSV 가져오기: Research → Test Suite → Catalog 가져오기")

    # 11
    add_h(doc, "11. Settings — 설정")
    add_table(
        doc,
        ["항목", "설명"],
        [
            ["언어 (Language)", "한국어 / English"],
            ["테마", "Dark / Light 모드"],
            ["Auto Refresh", "DB 상태 및 현재 화면 그리드 자동 갱신 (간격 설정)"],
            ["Analysis Web Worker", "Master Value 500자 이상일 때 분석을 워커 스레드에서 실행"],
            ["DB 백업 / 복원", "SQLite 파일 백업·덮어쓰기 복원"],
            ["Health Check", "코드·DB·엔진 상태 점검"],
            ["Packaging Verify", "배포 패키징 준비 상태 확인 (개발/QA용)"],
            ["메뉴 순서 초기화", "사이드바 레이아웃 리셋"],
        ],
    )
    add_p(doc, "설정 변경 후 「저장」 버튼을 눌러야 localStorage에 반영됩니다.")

    # 12
    add_h(doc, "12. 키보드 단축키")
    add_table(
        doc,
        ["단축키", "동작"],
        [
            ["Alt+1 ~ Alt+8", "사이드바 메뉴 순서대로 화면 이동"],
            ["Ctrl+Shift+I", "데이터 일괄 가져오기 모달"],
            ["F2", "신규 (Master, Code, CodeValue)"],
            ["Ctrl+S", "저장 (Master, Code, CodeValue)"],
            ["Delete", "삭제 (Master, Code, CodeValue)"],
        ],
    )
    add_p(doc, "입력 필드에 포커스가 있거나 확인 대화상자가 열려 있으면 CRUD 단축키는 동작하지 않습니다.")

    # 13
    add_h(doc, "13. 프로그램 정보")
    add_p(doc, "Toolbar의 「프로그램 정보」에서 앱 버전, DB 경로, 연결 상태, 로그 파일 경로를 확인합니다.")

    # 14
    add_h(doc, "14. 데이터 백업 및 복원")
    add_numbered(
        doc,
        [
            "Settings → 「DB 백업」: 현재 DB를 사용자가 지정한 경로에 저장.",
            "Settings → 「DB 복원」: 백업 파일 선택 후 확인 대화상자에서 승인.",
            "복원 후 앱 데이터가 갱신됩니다. 작업 전 백업을 권장합니다.",
        ],
    )

    # 15
    add_h(doc, "15. 문제 해결 (FAQ)")
    add_table(
        doc,
        ["증상", "조치"],
        [
            ["DB: 연결 안됨", "앱 재시작, Settings에서 DB 경로 확인, 로그 파일 검토"],
            ["저장 실패", "StatusBar 오류 메시지 확인 (유효성 검사·중복 번호 등)"],
            ["분석이 느림", "Settings에서 Web Worker 활성화 (대용량 Master)"],
            ["일괄 가져오기 실패", "파일 형식(TXT/CSV/XLSX)과 열 이름 확인"],
            ["언어가 섞여 보임", "Settings에서 언어 저장 후 화면 새로고침(메뉴 이동)"],
        ],
    )

    # 16
    add_h(doc, "16. 알려진 제한 (v1.0.0)")
    add_bullets(
        doc,
        [
            "Prediction / CodeValue: 구현됨, 레거시 1:1 미검증 (UI 배너 표시)",
            "P0-13 레거시 TestCase 카탈로그: 100건 DRAFT skeleton만 존재 (v1.1 예정)",
            "macOS / Linux 미지원",
            "Menu Edit: 사이드바 순서만 지원",
            "NSIS 설치본: 빌드는 완료, clean VM 수동 QA는 별도 진행",
        ],
    )

    # 17
    add_h(doc, "17. 참고 문서")
    add_bullets(
        doc,
        [
            "docs/STATUS.md — 구현·검증 현황",
            "docs/PRD.md — 제품 요구사항",
            "docs/RE-PLAYBOOK.md — 레거시 관측 → Verification 절차",
            "docs/INSTALL-QA-CHECKLIST.md — 설치 QA",
            "CHANGELOG.md — 릴리스 노트",
        ],
    )

    add_p(doc, "")
    footer = doc.add_paragraph("© 2026 CS E-Bid Team")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
